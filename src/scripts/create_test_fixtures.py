import os
import asyncio
from playwright.async_api import async_playwright
import docx

def create_fixtures():
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    fixture_dir = os.path.join(root, 'docs', '05_qa_and_demo', 'fixtures')
    os.makedirs(fixture_dir, exist_ok=True)

    # 1. vietnamese-knowledge.md (Rich 6-stage lifecycle & Hybrid RAG comparison)
    vietnamese_md = """# Chu Trình Quản Trị Tri Thức & Kiến Trúc Tìm Kiếm Lai KnowledgeOS

## 1. Chi Tiết 6 Bước Trong Chu Trình Quản Trị Tri Thức của KnowledgeOS

Hệ điều hành tri thức cá nhân KnowledgeOS vận hành theo chu trình khép kín gồm 6 giai đoạn liên tục:

1. **Thu thập dữ liệu (Collect)**: Tiếp nhận các nguồn tài liệu đa định dạng (PDF học thuật, Word giáo trình, Markdown ghi chú, bài giảng). Hệ thống tự động phân tách cấu trúc văn bản qua các Parser đa hình, chia nhỏ thành các đoạn 500 ký tự (độ gối đầu 100 ký tự) và lưu trữ tệp nhị phân gốc bền vững vào cơ sở dữ liệu.
2. **Tổ chức tri thức (Organize)**: Phân loại tài liệu khoa học theo cây Thư mục môn học (Collections) và Thẻ chuyên đề (Tags). Tích hợp thuật toán Trí tuệ Nhân tạo thông minh (Smart Organization Suggestions) tự động phân tích vector nội dung để gợi ý nhãn dán tối ưu.
3. **Thấu hiểu nội dung (Understand)**: Cung cấp không gian đọc Reader tập trung cao độ, loại bỏ hoàn toàn xao nhãng, hỗ trợ ghi chép nhanh các phát hiện nghiên cứu (Session Notes) và tự động phát hiện các tài liệu liên quan thông qua khoảng cách vector Cosine.
4. **Truy xuất lai (Retrieve)**: Kích hoạt đồng thời 2 nhánh tìm kiếm: Nhánh Ngữ nghĩa (Semantic Search qua pgvector với chỉ mục HNSW) và Nhánh Từ khóa chính xác (Lexical Search qua PostgreSQL Full-Text Search). Hợp nhất kết quả bằng thuật toán Xếp hạng Tương hỗ (Reciprocal Rank Fusion - RRF k=60).
5. **Hỏi đáp thông minh (Ask)**: Tổng hợp câu trả lời ngôn ngữ tự nhiên từ Google Gemini dựa trên ngữ cảnh tài liệu được cung cấp (Grounding), loại bỏ hoàn toàn hiện tượng AI bịa đặt ảo giác và gắn nhãn trích dẫn nguồn [1], [2] có thể bấm kiểm tra văn bản gốc.
6. **Học tập & Thống kê (Learn)**: Rèn luyện tính kỷ luật học tập với đồng hồ bấm giờ Pomodoro 25 phút không xao nhãng (Focus Mode) và theo dõi trực quan biểu đồ tăng trưởng tri thức thông qua bảng điều khiển Insights Dashboard.

---

## 2. Tại Sao Tìm Kiếm Lai (Hybrid RAG) Lại Vượt Trội Hơn Tìm Kiếm Vector Truyền Thống?

Tìm kiếm lai (Hybrid Retrieval) vượt trội hơn hẳn so với tìm kiếm vector đơn thuần nhờ các ưu thế công nghệ sau:

1. **Khắc phục điểm mù của Vector Embeddings**: Mô hình vector thường bị "trôi ngữ nghĩa" và không phân biệt được các mã định danh kỹ thuật, chuỗi số hiệu (như CVE-2026-8819, RFC-9421) hay tên biến camelCase. Nhánh Full-Text Search (FTS) với chỉ mục GIN của PostgreSQL xử lý hoàn hảo các chuỗi ký tự chính xác này.
2. **Hiểu sâu ngữ nghĩa câu hỏi tự nhiên**: Nhánh Semantic sử dụng vector embedding 768 chiều từ Gemini giúp hiểu được các câu hỏi đồng nghĩa, câu hỏi diễn giải gián tiếp ngay cả khi không trùng khớp từ vựng.
3. **Thuật toán Hợp nhất RRF (k=60) cân bằng tối ưu**: Không cần hiệu chỉnh trọng số thủ công, RRF tự động ưu tiên những đoạn văn xuất hiện ở thứ hạng cao trên cả hai nhánh tìm kiếm, mang lại độ chính xác vượt trội cho cả tiếng Việt và tiếng Anh.
"""
    with open(os.path.join(fixture_dir, 'vietnamese-knowledge.md'), 'w', encoding='utf-8') as f:
        f.write(vietnamese_md)

    # 2. exact-identifier.md (Cybersecurity CVE & Standards)
    exact_md = """# Báo Cáo An Ninh Mạng & Danh Mục Tiêu Chuẩn Kỹ Thuật

## 1. Thông Tin Chi Tiết Về Mã Lỗ Hổng CVE-2026-8819

- **Mã định danh lỗ hổng**: `CVE-2026-8819`
- **Tên lỗ hổng**: Lỗ hổng Thực thi Mã từ xa trong Bộ lọc Phiên làm việc (Spring Security Session Filter Remote Code Execution).
- **Điểm số nghiêm trọng CVSS**: **8.8 (Critical)**
- **Chuỗi Vector CVSS**: `CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H`
- **Mô tả kỹ thuật**: Lỗ hổng cho phép kẻ tấn công chưa xác thực gửi các tiêu đề HTTP Session được chế tạo đặc biệt nhằm làm tràn bộ đệm giải tuần tự hóa, từ đó chiếm quyền điều khiển tiến trình máy chủ.
- **Giải pháp khắc phục & Vá lỗi**:
  1. Nâng cấp ngay lập tức lên phiên bản Spring Boot 4.1.0 hoặc áp dụng bản vá bảo mật `patch-2026-08a`.
  2. Bật cờ cấu hình `server.servlet.session.cookie.same-site=strict` và kiểm tra chặt chẽ token CSRF trên toàn bộ endpoint POST/PUT.
  3. Kích hoạt bộ lọc xác thực dữ liệu đầu vào `StrictHeaderValidationFilter`.

---

## 2. Danh Mục Phần Cứng & Tiêu Chuẩn Quốc Tế

- **Tiêu chuẩn Ký số HTTP**: Tuân thủ nghiêm ngặt chuẩn `RFC-9421` về HTTP Message Signatures.
- **Tiêu chuẩn Token Web**: Tuân thủ chuẩn `RFC-7519` cho việc tuần tự hóa JSON Web Token (JWT).
- **Mã Bo mạch Chủ (Mainboard Serial)**: Mã định danh `KB-9902-REV4` đạt chứng nhận chịu nhiệt độ công nghiệp.
- **Vi xử lý Mật mã (Cryptographic Coprocessor)**: Chipset `HSM-AX771` hỗ trợ tăng tốc phần cứng cho các phép toán đường cong Elliptic (ECC).
- **Định danh Toàn cầu (Canonical UUID)**: `8f7b2c14-5d9a-4e8b-a2c3-9d1e4f5a6b7c`.
"""
    with open(os.path.join(fixture_dir, 'exact-identifier.md'), 'w', encoding='utf-8') as f:
        f.write(exact_md)

    # 3. oop-basics.md (Comprehensive OOP & Design Patterns)
    oop_md = """# Lập Trình Hướng Đối Tượng (OOP) & Các Mẫu Thiết Kế Phần Mềm

## 1. Bốn Trụ Cột của Lập Trình Hướng Đối Tượng

1. **Tính Đóng Gói (Encapsulation)**: Gom nhóm dữ liệu và các phương thức xử lý dữ liệu vào trong một đơn vị đối tượng, đồng thời che giấu trạng thái nội bộ bằng quyền truy cập `private`/`protected`. Trong KnowledgeOS, thực thể `Resource.java` bảo vệ máy trạng thái qua các hàm `beginParsing()`, `beginEmbedding()`, `markReady()`.
2. **Tính Kế Thừa (Inheritance)**: Cho phép một lớp con kế thừa lại các thuộc tính và hành vi của lớp cha nhằm tái sử dụng mã nguồn.
3. **Tính Đa Hình (Polymorphism)**: Khả năng các đối tượng khác nhau phản ứng với cùng một thông điệp theo các cách khác nhau. Thể hiện qua giao diện `ResourceParser` với các lớp triển khai `PdfResourceParser`, `DocxResourceParser`.
4. **Tính Trừu Tượng (Abstraction)**: Tập trung vào các hành vi cốt lõi của đối tượng và ẩn đi các chi tiết cài đặt phức tạp thông qua Interface hoặc Abstract Class.

---

## 2. Mẫu Thiết Kế Chiến Lược (Strategy Pattern) Trong RAG

Giao diện `RetrievalStrategy` định nghĩa phương thức tìm kiếm trừu tượng `retrieve()`. Có 3 chiến lược cụ thể:
- `SemanticRetrievalStrategy`: Tìm kiếm vector ngữ nghĩa với pgvector.
- `KeywordRetrievalStrategy`: Tìm kiếm từ khóa chính xác bằng PostgreSQL Full-Text Search.
- `HybridRetrievalStrategy`: Kết hợp cả hai nhánh và hợp nhất bằng Reciprocal Rank Fusion (RRF k=60).

Ưu điểm: Tuân thủ nguyên lý **Open/Closed Principle (OCP)**, cho phép mở rộng thuật toán mới mà không cần sửa đổi mã nguồn gọi dịch vụ.
"""
    with open(os.path.join(fixture_dir, 'oop-basics.md'), 'w', encoding='utf-8') as f:
        f.write(oop_md)

    # 4. medical-clinical-protocol.txt (Medical / Biomedical Domain)
    medical_txt = """BÁO CÁO ĐỀ CƯƠNG THỬ NGHIỆM LÂM SÀNG GIAI ĐOẠN III (PHASE III CLINICAL PROTOCOL)

1. Tên Hoạt Chất Nghiên Cứu: Cardiotrex-900 (Mã hợp chất thử nghiệm: CTX-900).
2. Chỉ Định Điều Trị: Điều trị suy tim sung huyết mạn tính có phân suất tống máu giảm (HFrEF).
3. Liều Lượng & Phác Đồ: 50mg/ngày, uống 1 lần duy nhất vào buổi sáng sau bữa ăn.
4. Quy Mô Mẫu Thử Nghiệm: 1.500 bệnh nhân tại 42 trung tâm y khoa quốc tế.
5. Hiệu Quả Lâm Sàng (Primary Efficacy Endpoint):
   - Tỷ lệ giảm biến cố tim mạch nhập viện đạt 94.2% so với nhóm dùng giả dược (Placebo).
   - Cải thiện chỉ số phân suất tống máu thất trái (LVEF) trung bình tăng thêm 8.5% sau 24 tuần.
6. Tác Dụng Không Mong Muốn (Adverse Events):
   - Tỷ lệ chóng mặt nhẹ: 1.8%
   - Tỷ lệ hạ huyết áp tư thế: 0.9%
   - Không ghi nhận trường hợp độc tính trên gan hay suy giảm chức năng thận cấp.
"""
    with open(os.path.join(fixture_dir, 'medical-clinical-protocol.txt'), 'w', encoding='utf-8') as f:
        f.write(medical_txt)

    # 5. macroeconomic-outlook-2026.txt (Economics & Finance Domain)
    finance_txt = """BÁO CÁO TRIỂN VỌNG KINH TẾ VĨ MÔ & TÀI CHÍNH NĂM 2026

1. Tăng Trưởng Tổng Sản Phẩm Quốc Nội (GDP):
   - Dự báo mức tăng trưởng kinh tế toàn năm 2026 đạt 6.8%, dẫn đầu bởi khu vực sản xuất chế biến và xuất khẩu điện tử.
2. Kiểm Soát Lạm Phát & Chỉ Số Giá Tiêu Dùng (CPI):
   - Lạm phát mục tiêu được kiểm soát chặt chẽ ở mức 3.2%, thấp hơn trần định hướng 4.0% của Ngân hàng Trung ương.
3. Chính Sách Tiền Tệ & Lãi Suất:
   - Lãi suất tái cấp vốn duy trì ở mức 4.5%/năm.
   - Dự kiến có 1 đợt cắt giảm lãi suất 25 điểm cơ bản (bps) vào Quý 3 năm 2026 để kích cầu đầu tư tư nhân.
4. Vốn Đầu Tư Trực Tiếp Nước Ngoài (FDI):
   - Tổng vốn FDI đăng ký đạt 28.5 tỷ USD, trong đó 65% tập trung vào công nghiệp bán dẫn và trung tâm dữ liệu AI.
"""
    with open(os.path.join(fixture_dir, 'macroeconomic-outlook-2026.txt'), 'w', encoding='utf-8') as f:
        f.write(finance_txt)

    # 6. Generate DOCX file: software-engineering-guide.docx
    doc = docx.Document()
    doc.add_heading('Giáo Trình Kỹ Nghệ Phần Mềm & Kiến Trúc Hệ Thống', 0)
    doc.add_paragraph('Tài liệu hướng dẫn xây dựng hệ thống phần mềm doanh nghiệp hiện đại.')
    doc.add_heading('1. Nguyên Lý Clean Architecture', level=1)
    doc.add_paragraph('Clean Architecture phân tách hệ thống thành các vòng tròn đồng tâm, trong đó quy tắc phụ thuộc chỉ được phép trỏ từ ngoài vào trong. Lớp thực thể (Entities) nằm ở trung tâm và không phụ thuộc vào bất kỳ framework hay cơ sở dữ liệu nào.')
    doc.add_heading('2. Đảm Bảo An Toàn Giao Dịch ACID', level=1)
    doc.add_paragraph('Mọi thao tác cập nhật dữ liệu quan trọng đều phải được thực thi trong một giao dịch duy nhất có tính chất Nguyên tử (Atomicity), Nhất quán (Consistency), Cô lập (Isolation) và Bền vững (Durability).')
    docx_path = os.path.join(fixture_dir, 'software-engineering-guide.docx')
    doc.save(docx_path)
    print(f"Generated DOCX fixture: {docx_path}")

    # Also update test files in refer/qa_dataset/fixtures/ so automated tests are synced
    refer_qa_dir = os.path.join(root, 'refer', 'qa_dataset', 'fixtures')
    os.makedirs(refer_qa_dir, exist_ok=True)
    with open(os.path.join(refer_qa_dir, 'vietnamese-knowledge.md'), 'w', encoding='utf-8') as f:
        f.write(vietnamese_md)
    with open(os.path.join(refer_qa_dir, 'exact-identifier.md'), 'w', encoding='utf-8') as f:
        f.write(exact_md)

async def generate_pdf_fixture():
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    fixture_dir = os.path.join(root, 'docs', '05_qa_and_demo', 'fixtures')
    pdf_path = os.path.join(fixture_dir, 'ai-security-handbook.pdf')
    html_temp = os.path.join(fixture_dir, 'ai-security-handbook.html')

    html_content = """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>Sổ Tay An Toàn Trí Tuệ Nhân Tạo & Phòng Thủ Prompt Injection</title>
<style>
body { font-family: "Times New Roman", serif; padding: 30px; font-size: 14pt; line-height: 1.6; }
h1 { color: #0f2d4a; font-size: 22pt; }
h2 { color: #1a4971; font-size: 16pt; margin-top: 20px; }
.badge { background: #e2e8f0; padding: 3px 8px; border-radius: 4px; font-family: monospace; }
</style>
</head>
<body>
<h1>Sổ Tay An Toàn Trí Tuệ Nhân Tạo & Phòng Thủ Prompt Injection</h1>
<p><em>Tài liệu nghiên cứu kỹ thuật bảo mật ứng dụng LLM và hệ thống Retrieval-Augmented Generation.</em></p>

<h2>1. Tổng Quan Về Tấn Công Prompt Injection</h2>
<p>Tấn công Prompt Injection xảy ra khi kẻ tấn công chèn các câu lệnh điều khiển hệ thống vào trong dữ liệu người dùng (tệp PDF, văn bản tải lên) nhằm làm sai lệch hành vi của mô hình ngôn ngữ lớn (LLM).</p>

<h2>2. Cơ Chế Phòng Thủ Đa Tầng Của KnowledgeOS</h2>
<p>1. <strong>Cách ly Dữ liệu bằng Thẻ XML Thụ Động</strong>: Lớp <span class="badge">GroundedPromptBuilder</span> bọc toàn bộ nội dung trích xuất vào các thẻ XML &lt;evidence&gt; và chỉ định rõ đây là dữ liệu tham khảo, tuyệt đối không được coi là câu lệnh điều khiển.</p>
<p>2. <strong>Chỉ Thị Nghiêm Ngặt Chống Ảo Giác</strong>: Hệ thống bắt buộc AI từ chối trả lời nếu tài liệu không chứa đủ dữ liệu, loại bỏ hoàn toàn hiện tượng suy đoán vô căn cứ.</p>
<p>3. <strong>Bảo Mật Bộ Nhớ Đệm và Phiên</strong>: Mã hóa toàn bộ dữ liệu nhị phân trong bảng <span class="badge">storage_blobs</span> với cơ chế phân quyền <span class="badge">owner_id</span> tuyệt đối.</p>
</body>
</html>
"""
    with open(html_temp, 'w', encoding='utf-8') as f:
        f.write(html_content)

    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()
        await page.goto(f"file:///{html_temp.replace(os.sep, '/')}")
        await page.pdf(path=pdf_path, format="A4", margin={"top": "20mm", "bottom": "20mm", "left": "20mm", "right": "20mm"})
        await browser.close()
    
    if os.path.exists(html_temp):
        os.remove(html_temp)
    print(f"Generated PDF fixture: {pdf_path}")

if __name__ == '__main__':
    create_fixtures()
    asyncio.run(generate_pdf_fixture())
